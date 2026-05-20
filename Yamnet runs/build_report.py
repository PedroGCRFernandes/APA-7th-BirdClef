"""Generate the YAMNet / Gemma4 approach report as a Word document."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)


def H1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)


def H2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)


def H3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.italic = True
    r.font.size = Pt(11.5)


def P(text):
    doc.add_paragraph(text)


def B(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def CODE(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9.5)


# ─── Title ────────────────────────────────────────────────────────────────────
H1("BirdCLEF 2026 — YAMNet + Gemma4 Agent")
sub = doc.add_paragraph()
r = sub.add_run("Track B Autonomous Research Agent · APA 7th coursework report")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# ─── 1. First approach ────────────────────────────────────────────────────────
H2("1. First Approach")
P(
    "The original plan for Track B was an end-to-end CNN trained from scratch on "
    "mel-spectrograms of the BirdCLEF 2026 training clips. An EfficientNet-B0 backbone "
    "was wired into the agent loop, with the LLM only choosing classification-head "
    "topology. On an Apple-Silicon M1 Pro (16 GB) this proved unworkable: tensorflow-metal "
    "kernels for depthwise convolutions and the Swish activation in EfficientNet were "
    "pathologically slow — often slower than CPU — and a single training epoch on the "
    "full clip set took longer than the entire iteration budget allowed."
)
P(
    "The pivot was to treat the audio backbone as a frozen feature extractor and let "
    "the LLM agent design only the head sitting on top of pre-computed embeddings. "
    "YAMNet (Google AudioSet, MobileNet-v1 trunk) was chosen for three reasons: (a) it "
    "ingests raw 16 kHz mono waveforms directly, removing the spectrogram pipeline, "
    "(b) its 1024-dimensional embedding is well-studied as a transfer-learning feature "
    "for bioacoustics, and (c) the SavedModel is small enough to bundle into the "
    "Kaggle submission notebook for fully offline inference."
)
P(
    "The local research loop runs entirely on CPU; the LLM that proposes head architectures "
    "is Gemma 4 (e4b variant) served by Ollama. Gemma was chosen over larger models because "
    "it returns a fresh proposal in seconds, keeping the iteration cadence high enough that "
    "20 rounds complete overnight without GPU access."
)

# ─── 2. System architecture ───────────────────────────────────────────────────
H2("2. System Architecture")
P(
    "The agent has three decoupled stages: data preparation (run once), the LLM design "
    "loop (run per iteration), and Kaggle-bound model export (run when a new global best "
    "is reached). All stages live in agent.py and share an embedding cache held in RAM."
)

H3("Data pipeline")
B([
    "load_audio() decodes 5 s of mono 16 kHz audio per clip using librosa; clips shorter "
    "than 5 s are zero-padded.",
    "audio_to_embedding() calls the frozen YAMNet SavedModel once per waveform and returns "
    "the full frame-level output of shape (N_FRAMES, 1024). The frame count is probed at "
    "startup on a silence clip and cached.",
    "precompute_embeddings() runs decode in a ThreadPoolExecutor (librosa releases the "
    "GIL) and YAMNet inference on the main thread (the SavedModel is not thread-safe).",
    "Per clip, the cache stores 1 clean embedding plus N_AUG_VARIANTS (=6) acoustic "
    "variants produced by augment_waveform (pitch shift, time stretch, Gaussian field "
    "noise, random gain, time shift). Training picks a variant at random each epoch; "
    "validation always uses the clean variant.",
])

H3("LLM agent loop")
B([
    "build_prompt() in prompt_builder.py composes a prompt that fixes the input shape "
    "and mandatory boilerplate, lists the available Keras layers, and embeds the full "
    "history of previous iterations (architecture, val_auc, val_loss, errors).",
    "If the LLM's code crashes, ask_llm_to_fix() retries up to MAX_FIX_RETRIES (=3) "
    "times with the error message and a stricter contract reminder.",
    "After each successful run, ask_llm_to_analyze() asks the LLM to reflect on the "
    "result; its summary is stored in the experiment log for the next prompt's history.",
    "A stagnation detector (_detect_stagnation) flips the prompt into 'structural change' "
    "mode once val_auc plateaus, nudging the LLM toward focal loss, multi-branch heads, "
    "self-attention, or different LR/regularisation regimes.",
])

H3("Training stage")
B([
    "execute_safely() runs the LLM-emitted code in a restricted namespace that exposes "
    "only the whitelisted layer set and the frozen backbone_model.",
    "model.compile is forced here (not by the LLM) with AdamW + cosine-decay LR + "
    "BinaryCrossentropy(label_smoothing=0.02), and macro-AUROC as the validation metric.",
    "Two callbacks control training: _BestWeightsCallback keeps the best val_auc weights "
    "in RAM to avoid serialisation churn, and PlateauCallback stops early when val_auc "
    "stops moving by at least 0.005 absolute over 5 epochs.",
    "Soundscape validation (SOUNDSCAPE_VAL=True) mixes held-out soundscape windows into "
    "both training and validation from epoch 1, aligning the val_auc proxy with the real "
    "Kaggle objective rather than clean-clip AUC.",
])

# ─── 3. Code characteristics ──────────────────────────────────────────────────
H2("3. Characteristics of the Code Base")
B([
    "Single-file orchestration (agent.py, ~1100 lines) with two small helpers "
    "(prompt_builder.py, experiment_log.py). All configuration is module-level constants "
    "at the top, no YAML or argparse layer.",
    "Defensive Kaggle-compatibility: a _check_versions() guard at startup aborts if "
    "tensorflow/keras versions diverge from the Kaggle BirdCLEF image, and "
    "save_keras_clean() strips the dead renorm_* config keys that newer Keras writes but "
    "Kaggle's older Keras refuses to load.",
    "Two-tier best-model tracking: each run keeps its own runs/<id>/best_model.keras, but "
    "the canonical models/best_model.keras is only overwritten when the new run beats the "
    "all-time best val_auc on disk — protecting the Kaggle copy from regressions.",
    "Three layers of augmentation: waveform-level acoustic variants in the cache, "
    "embedding-axis band masking (SpecAugment analogue, EMB_MASK_WIDTH=128), and "
    "batch-level mixup with Beta(0.4, 0.4) blending applied inside the generator.",
    "Hardware-adaptive parallelism: _DATA_WORKERS = cpu_count − 2 so two cores stay free "
    "for TensorFlow and the OS; multiprocessing is disabled deliberately to keep the "
    "shared embedding cache from being copied per worker.",
    "Honest sanity-checking: _auc_sanity_warnings prints loud warnings when epoch-1 "
    "val_auc looks too high (label/validation leakage) or too low (broken taxonomy map).",
])

# ─── 4. Debugging journey ─────────────────────────────────────────────────────
H2("4. Debugging Journey")
P(
    "Most of the engineering time was spent not on the head designs themselves but on "
    "fighting the gap between what a 4-billion-parameter local LLM can reliably produce "
    "and what Keras 3 will accept. The notable issues are summarised below."
)

H3("4.1 Silent label-mapping leakage")
P(
    "Early runs showed first-epoch val_auc of ~0.99, which was suspicious for a "
    "234-class multi-label task. The cause was a type mismatch: taxonomy primary_label "
    "values were read as a mix of strings and integers by pandas, but soundscape labels "
    "were cast to int(taxon_id). About 60% of soundscape windows then mapped to no class "
    "and silently became all-zero target vectors, inflating micro-AUC. Fixed by "
    "normalising every key to a stripped string at lookup time and dropping any window "
    "whose label vector sums to zero."
)

H3("4.2 Wrong AUC flavour")
P(
    "The default tf.keras.metrics.AUC computes micro-AUROC over the flattened "
    "(samples × classes) tensor. With 1–2 positive labels per row out of 234, that "
    "metric is trivially high. The agent now compiles with "
    "AUC(multi_label=True, num_labels=234), matching Kaggle's macro-averaged scoring."
)

H3("4.3 Kaggle / local Keras divergence")
P(
    "A model trained locally on Keras 3.14 failed to load on the Kaggle BirdCLEF image "
    "(Keras 3.10): the local Keras serialises BatchNormalization with renorm, "
    "renorm_clipping, and renorm_momentum keys that the older Keras rejects. "
    "save_keras_clean() now opens the .keras zip after saving, strips those keys from "
    "config.json, and writes the archive back."
)

H3("4.4 LLM contract failures")
P(
    "Gemma 4 (e4b) breaks the agent contract in three recurring ways: (a) it omits the "
    "mandatory call to backbone_model entirely, (b) it produces a uniformly-indented "
    "snippet that fails Python's exec() on line 1, and (c) it passes the backbone Model "
    "object as a positional argument to a layer call, which Keras 3 explicitly forbids. "
    "The fixes were to mark the boilerplate as MANDATORY (not example) in the prompt, "
    "to textwrap.dedent the extracted code before exec, and to surface the "
    "'no positional Model args' rule in the fix-retry prompt with the exact input shape."
)

H3("4.5 Apple-Silicon GPU detour")
P(
    "tensorflow-metal was tried and discarded. Depthwise convolutions and Swish on Metal "
    "ran several times slower than the CPU path. The default is now USE_GPU=False, with "
    "a comment explaining the trade-off in case a future backbone benefits from Metal."
)

# ─── 5. Limitations ───────────────────────────────────────────────────────────
H2("5. Limitations")
B([
    "Frozen backbone ceiling: YAMNet was trained on AudioSet, not on fine-grained bird "
    "vocalisations. Whatever inter-species discrimination is not already present in the "
    "1024-d embedding cannot be recovered by any head, no matter how clever.",
    "Local LLM ceiling: Gemma 4 (e4b) consumes ~20% of each iteration in crash-and-fix "
    "retries, and its proposals cluster around shallow variations on Dense + Dropout + "
    "GlobalAveragePooling1D. The plateau-mode prompt mitigates this but does not solve it.",
    "Single-fold validation: a single 80/20 soundscape split is small and noisy. val_auc "
    "differences below ~0.01 between runs are within the noise floor of the val set and "
    "cannot reliably rank head architectures.",
    "No cross-recording leakage check: train/val splits are stratified by primary_label "
    "but not grouped by recordist or geographic location. Some apparent generalisation "
    "may be memorisation of recordist-specific background noise.",
    "RAM-bound cache: with N_AUG_VARIANTS=6 and the full clip set, the embedding cache "
    "approaches 12 GB. On a 16 GB machine this leaves little headroom for TF runtime; "
    "the agent prints a warning above 12 GB but does not spill to disk.",
    "Single-model inference: the Kaggle submission uses the single best_model.keras "
    "rather than an ensemble. Test-time augmentation is also disabled.",
])

# ─── 6. Future work ───────────────────────────────────────────────────────────
H2("6. Future Changes")
B([
    "Swap the LLM. Even Qwen 2.5 7B would dramatically reduce the crash-and-fix overhead, "
    "and a cloud-hosted Sonnet or GPT-4-class model would unlock substantially richer "
    "head designs without changing any other code.",
    "Auto-inject the mandatory boilerplate. Prepend the `inputs = Input(...); "
    "x = backbone_model(inputs, training=False)` lines automatically and ask the LLM only "
    "for the segment starting at the first transform — eliminating contract-failure "
    "classes (a) and (c) entirely.",
    "Add a second frozen backbone (e.g. BirdNET or PANNs CNN14) and let the head consume "
    "the concatenation. BirdNET is taxonomy-aligned with this competition and would "
    "almost certainly lift macro-AUROC above what YAMNet alone can reach.",
    "Replace single-fold val with k-fold (k=3) recordist-grouped cross-validation. "
    "Slower per iteration, but the val_auc signal becomes trustworthy enough to compare "
    "runs that differ by <1 percentage point.",
    "Run-level ensembling. The agent already stores every run's best_model.keras; "
    "loading the top-N by val_auc and averaging their sigmoid outputs at submission "
    "time costs nothing locally and reliably reduces variance on Kaggle.",
    "Add test-time augmentation: average predictions over 3–5 random 5-s windows per "
    "soundscape rather than a single fixed window.",
    "Persist embedding cache to disk (np.memmap or .npz). Currently a crashed run "
    "discards the ~12 GB cache and the next start re-decodes every clip.",
])

# ─── 7. Reflection ────────────────────────────────────────────────────────────
H2("7. Reflection on the Agent Paradigm")
P(
    "The autonomous-research framing imposes a specific cost structure: most of the time "
    "is spent making the LLM's output executable, not improving the model. With a strong "
    "LLM that cost is small and the loop becomes a genuine search over architectures. "
    "With a 4-billion-parameter local model the loop is closer to constrained random "
    "search with English-language hints — useful as scaffolding, but not a substitute for "
    "a stronger backbone or a larger LLM."
)
P(
    "The clearest win of this project is the separation of concerns: data, model, and "
    "search are independent. Replacing YAMNet with BirdNET, or Gemma with Sonnet, or the "
    "single-fold val with k-fold, are all one-file changes. That decoupling is the "
    "deliverable I would re-use, more than any specific val_auc number."
)

out = "/Users/ricardodesiderio/Desktop/APA-7th-BirdClef/YAMNet_Gemma4_Report.docx"
doc.save(out)
print(f"Saved: {out}")
